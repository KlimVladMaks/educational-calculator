import json
import datetime
from docx import Document
from openpyxl import Workbook


class Importer:
    """
    Класс для импорта данных об учебных группах из JSON файла (database.json)
    и экспорта их в форматы DOCX и XLSX.

    """
    def __init__(self, json_file_path):
        self.json_file_path = json_file_path
        self.data = self.load_data()
        # Создадим словарь календарей по их названию для быстрого доступа
        self.calendars = {cal["name"]: cal for cal in self.data.get("calendars", [])}
        # Список групп (учебные группы)
        self.groups = self.data.get("groups", [])

    def load_data(self):
        """
        Загружает данные из JSON файла.

        """
        try:
            with open(self.json_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, dict):
                    return data
                else:
                    print("Ошибка: JSON-файл должен содержать объект с ключами 'calendars', 'programs' и 'groups'.")
                    return {}
        except FileNotFoundError:
            print(f"Файл {self.json_file_path} не найден.")
            return {}
        except json.JSONDecodeError as e:
            print("Ошибка декодирования JSON:", e)
            return {}

    def compute_group_metrics(self, group):
        """
        Для заданной учебной группы вычисляет:
          - Дата окончания обучения (берется из календаря)
          - Общее время обучения (в днях)
          - Число неучебных дней (на основе days_off_list календаря)
          - Число учебных дней (общие дни минус неучебные)

        """
        group_start_str = group.get("start_date", "")
        calendar_name = group.get("calendar", "")
        if not group_start_str or calendar_name not in self.calendars:
            print(f"Ошибка: Для группы '{group.get('name', '')}' отсутствует дата начала или не найден календарь '{calendar_name}'.")
            return None, None, None, None, None

        try:
            group_start = datetime.datetime.strptime(group_start_str, "%Y-%m-%d")
        except ValueError as e:
            print(f"Ошибка преобразования даты для группы '{group.get('name', '')}':", e)
            return None, None, None, None, None

        calendar = self.calendars[calendar_name]
        calendar_end_str = calendar.get("end_date", "")
        try:
            calendar_end = datetime.datetime.strptime(calendar_end_str, "%Y-%m-%d")
        except ValueError as e:
            print(f"Ошибка преобразования даты окончания календаря '{calendar_name}':", e)
            return None, None, None, None, None

        # Дата окончания обучения группы берется из календаря
        end_date_str = calendar_end_str

        # Общее время обучения: число дней между group.start_date и calendar.end_date (включительно)
        total_days = (calendar_end - group_start).days + 1
        if total_days < 0:
            print(f"Ошибка: Дата начала группы '{group.get('name', '')}' позже даты окончания календаря '{calendar_name}'.")
            return None, None, None, None, None

        # Подсчет неучебных дней: считаем, сколько дней из days_off_list попадают в период [group_start, calendar_end]
        non_training_days = 0
        for day_str in calendar.get("days_off_list", []):
            try:
                day = datetime.datetime.strptime(day_str, "%Y-%m-%d")
                if group_start <= day <= calendar_end:
                    non_training_days += 1
            except ValueError:
                continue

        training_days = total_days - non_training_days

        return group_start_str, end_date_str, training_days, non_training_days, total_days

    def import_to_word(self, docx_file_path):
        """
        Экспортирует данные учебных групп в формате DOCX.
        """
        doc = Document()
        doc.add_heading("Учебные группы", level=1)

        headers = [
            "Название", "Календарь", "Программа", 
            "Дата начала обучения", "Дата окончания обучения", 
            "Число учебных дней", "Число неучебных дней", "Общее время обучения (в днях)"
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        for group in self.groups:
            # Получаем вычисленные показатели для группы
            start_date, end_date, training_days, non_training_days, total_days = self.compute_group_metrics(group)
            # Если вычисления не прошли, пропускаем группу
            if start_date is None:
                continue

            row_cells = table.add_row().cells
            row_cells[0].text = group.get("name", "")
            row_cells[1].text = group.get("calendar", "")
            row_cells[2].text = group.get("program", "")
            row_cells[3].text = start_date
            row_cells[4].text = end_date
            row_cells[5].text = str(training_days)
            row_cells[6].text = str(non_training_days)
            row_cells[7].text = str(total_days)

        try:
            doc.save(docx_file_path)
            print(f"Данные успешно экспортированы в DOCX файл: {docx_file_path}")
        except Exception as e:
            print("Ошибка сохранения DOCX файла:", e)

    def import_to_excel(self, excel_file_path):
        """
        Экспортирует данные учебных групп в формате XLSX.
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Учебные группы"

        headers = [
            "Название", "Календарь", "Программа",
            "Дата начала обучения", "Дата окончания обучения",
            "Число учебных дней", "Число неучебных дней", "Общее время обучения (в днях)"
        ]
        ws.append(headers)

        for group in self.groups:
            start_date, end_date, training_days, non_training_days, total_days = self.compute_group_metrics(group)
            if start_date is None:
                continue
            row = [
                group.get("name", ""),
                group.get("calendar", ""),
                group.get("program", ""),
                start_date,
                end_date,
                training_days,
                non_training_days,
                total_days
            ]
            ws.append(row)

        try:
            wb.save(excel_file_path)
            print(f"Данные успешно экспортированы в Excel файл: {excel_file_path}")
        except Exception as e:
            print("Ошибка сохранения Excel файла:", e)


# Пример использования класса Importer с файлом database.json
if __name__ == "__main__":
    importer = Importer("database.json")
    importer.import_to_word("educational_groups.docx")
    importer.import_to_excel("educational_groups.xlsx")
